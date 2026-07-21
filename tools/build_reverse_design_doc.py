from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "documents" / "항해의_끝_해적왕의_유산_역기획서.docx"
ASSET_DIR = ROOT / "tmp" / "doc_assets"

# compact_reference_guide token map with a named game-brand accent override.
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
CONTENT_WIDTH_IN = 6.5
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BODY_FONT = "Noto Sans KR"
EAST_ASIA_FONT = "Noto Sans KR"
BODY_SIZE = 11
BODY_AFTER = 6
BODY_LINE = 1.25
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
INK = "202A33"
MUTED = "687782"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BRASS = "B8862B"
SEA = "176071"
RISK = "9B1C1C"
POSITIVE = "1F3A5F"

FONT_REGULAR = Path("/System/Library/Fonts/AppleSDGothicNeo.ttc")
FONT_BOLD = Path("/Users/2jh0926/Library/Fonts/NotoSansCJKkr-Bold.otf")


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, margins: dict[str, int] = CELL_MARGINS_DXA) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in margins.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D8DEE5", size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None,
                 name: str = BODY_FONT) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def configure_style(style, size: float, color: str, bold: bool,
                    before: float, after: float, line: float = BODY_LINE) -> None:
    style.font.name = BODY_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = rgb(color)
    r_pr = style._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), BODY_FONT)
    r_pr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    r_pr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = True
    pf.widow_control = True


def add_numbering_definition(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else f"%{level + 1}.")
        lvl.append(lvl_text)
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "tab")
        lvl.append(suffix)

        base_left = 540 + level * 360
        hanging = 270
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(base_left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(base_left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def shade_paragraph(paragraph, fill: str, left_color: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if left_color:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_color)
        borders.append(left)
        p_pr.append(borders)


def set_paragraph_padding(paragraph, before: float = 6, after: float = 6,
                          left: float = 8, right: float = 8) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.left_indent = Pt(left)
    pf.right_indent = Pt(right)


def add_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True, color=INK)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, color=INK)
    return p


def add_bullet(doc: Document, text: str, bullet_num_id: int, level: int = 0):
    p = doc.add_paragraph(style="Normal")
    apply_numbering(p, bullet_num_id, level)
    set_run_font(p.add_run(text), color=INK)
    return p


def add_number(doc: Document, text: str, decimal_num_id: int, level: int = 0):
    p = doc.add_paragraph(style="Normal")
    apply_numbering(p, decimal_num_id, level)
    set_run_font(p.add_run(text), color=INK)
    return p


def add_callout(doc: Document, label: str, text: str, color: str = POSITIVE,
                fill: str = CALLOUT):
    p = doc.add_paragraph()
    shade_paragraph(p, fill, color)
    set_paragraph_padding(p, before=7, after=7, left=10, right=8)
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, size=10.5, bold=True, color=color)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    return p


def add_formula(doc: Document, formula: str, explanation: str | None = None):
    p = doc.add_paragraph()
    shade_paragraph(p, "EEF3F6", DARK_BLUE)
    set_paragraph_padding(p, before=6, after=6, left=10, right=8)
    run = p.add_run(formula)
    set_run_font(run, name=BODY_FONT, size=9.2, bold=True, color=NAVY)
    if explanation:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Pt(10)
        p2.paragraph_format.space_after = Pt(7)
        set_run_font(p2.add_run(explanation), size=9.3, italic=True, color=MUTED)
    return p


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths_dxa: Sequence[int], font_size: float = 9.2,
              header_fill: str = PALE_BLUE, alignments: Sequence[str] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for index, header in enumerate(headers):
        cell = hdr.cells[index]
        set_cell_fill(cell, header_fill)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(str(header)), size=font_size, bold=True, color=NAVY)

    for row_values in rows:
        row = table.add_row()
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.13
            if alignments and alignments[index] == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(str(value)), size=font_size, color=INK)

    set_table_geometry(table, widths_dxa)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(text), size=9, italic=True, color=MUTED)
    return p


def set_picture_alt(paragraph, title: str, description: str) -> None:
    for doc_pr in paragraph._p.xpath(".//wp:docPr"):
        doc_pr.set("title", title)
        doc_pr.set("descr", description)


def heading(doc: Document, level: int, text: str):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), bold=True, color=BLUE if level < 3 else DARK_BLUE,
                 size={1: 16, 2: 13, 3: 12}[level])
    return p


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def draw_text_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str,
                     font: ImageFont.FreeTypeFont, fill: str, line_gap: int = 8) -> None:
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 28
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y1 + ((y2 - y1) - total) / 2
    for line, height in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        draw.text((x1 + ((x2 - x1) - width) / 2, y), line, font=font, fill=fill)
        y += height + line_gap


def font(size: int, bold: bool = False):
    path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    return ImageFont.truetype(str(path), size=size)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int],
          color: str = "#2E74B5", width: int = 5) -> None:
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    dx, dy = ex - sx, ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    base_x, base_y = ex - ux * 18, ey - uy * 18
    points = [
        (ex, ey),
        (base_x + px * 9, base_y + py * 9),
        (base_x - px * 9, base_y - py * 9),
    ]
    draw.polygon(points, fill=color)


def make_core_loop(path: Path) -> None:
    image = Image.new("RGB", (1500, 760), "#F6F8FA")
    draw = ImageDraw.Draw(image)
    title_font = font(34, True)
    body_font = font(24, True)
    small_font = font(18)
    draw.text((60, 40), "한 번의 항해와 전승 루프", font=title_font, fill="#0B2545")
    boxes = [
        (60, 150, 300, 300, "거점 섬", "전승 강화\n선장 선택"),
        (360, 150, 600, 300, "출항", "초기 자원\n선원 1명"),
        (660, 150, 900, 300, "항로 선택", "보급 소모\n노드 진입"),
        (960, 150, 1200, 300, "노드 해결", "전투·항구\n조우·보물"),
        (960, 440, 1200, 590, "보상과 성장", "금화·악명\n선원·유물"),
        (660, 440, 900, 590, "보스 / 다음 Act", "3개 해역\n최종 보물"),
        (270, 440, 570, 590, "항해 종료", "선체 0 또는 사기 0\n승리 / 패배"),
    ]
    for x1, y1, x2, y2, label, sub in boxes:
        draw.rectangle((x1, y1, x2, y2), fill="#FFFFFF", outline="#6F8794", width=3)
        draw_text_center(draw, (x1, y1 + 10, x2, y1 + 76), label, body_font, "#1F4D78")
        draw_text_center(draw, (x1, y1 + 72, x2, y2 - 8), sub, small_font, "#526572")
    arrow(draw, (300, 225), (360, 225))
    arrow(draw, (600, 225), (660, 225))
    arrow(draw, (900, 225), (960, 225))
    arrow(draw, (1080, 300), (1080, 440))
    arrow(draw, (960, 515), (900, 515))
    arrow(draw, (780, 440), (780, 300), color="#176071")
    arrow(draw, (660, 515), (570, 515), color="#9B1C1C")
    arrow(draw, (270, 515), (175, 515))
    draw.line([(175, 515), (175, 300)], fill="#B8862B", width=5)
    arrow(draw, (175, 300), (175, 300), color="#B8862B")
    draw.text((1215, 475), "일반 노드는\n항로 선택으로 복귀", font=small_font, fill="#526572")
    draw.text((65, 650), "런 안에서는 빌드를 누적하고, 런 밖에서는 획득 악명을 영구 능력치로 전환한다.", font=font(21), fill="#384A55")
    image.save(path)


def make_state_flow(path: Path) -> None:
    image = Image.new("RGB", (1500, 900), "#F6F8FA")
    draw = ImageDraw.Draw(image)
    draw.text((60, 38), "런타임 상태 전이", font=font(34, True), fill="#0B2545")
    nodes = {
        "harbor": (60, 160, 260, 260, "거점 / run 없음"),
        "intro": (340, 160, 560, 260, "interstitial"),
        "map": (650, 160, 850, 260, "map"),
        "resolve": (960, 160, 1160, 260, "resolving"),
        "event": (930, 390, 1110, 490, "event"),
        "port": (1160, 390, 1340, 490, "port"),
        "combat": (650, 390, 850, 490, "combat"),
        "reward": (650, 620, 850, 720, "reward"),
        "gameover": (280, 620, 500, 720, "gameover"),
        "victory": (980, 620, 1180, 720, "victory"),
    }
    for key, (x1, y1, x2, y2, label) in nodes.items():
        fill = "#FFF7E6" if key in {"harbor", "victory"} else "#FFFFFF"
        outline = "#B8862B" if key in {"harbor", "victory"} else "#6F8794"
        draw.rectangle((x1, y1, x2, y2), fill=fill, outline=outline, width=3)
        draw_text_center(draw, (x1, y1, x2, y2), label, font(22, True), "#1F4D78")
    arrow(draw, (260, 210), (340, 210))
    arrow(draw, (560, 210), (650, 210))
    arrow(draw, (850, 210), (960, 210))
    arrow(draw, (1060, 260), (1020, 390))
    arrow(draw, (1115, 260), (1240, 390))
    arrow(draw, (960, 235), (850, 430))
    arrow(draw, (750, 490), (750, 620))
    arrow(draw, (930, 440), (850, 235), color="#176071")
    arrow(draw, (1160, 440), (850, 235), color="#176071")
    arrow(draw, (650, 670), (500, 670), color="#9B1C1C")
    arrow(draw, (850, 670), (980, 670), color="#B8862B")
    draw.line([(750, 620), (620, 560), (620, 260)], fill="#176071", width=5)
    arrow(draw, (620, 260), (650, 235), color="#176071")
    draw.line([(390, 620), (390, 545), (110, 545), (110, 260)], fill="#9B1C1C", width=5)
    arrow(draw, (110, 260), (110, 260), color="#9B1C1C")
    draw.line([(1080, 720), (1080, 800), (110, 800), (110, 260)], fill="#B8862B", width=5)
    draw.text((70, 825), "공통 실패 조건: 선체 ≤ 0 또는 사기 ≤ 0. 포기는 보상 은행 처리 없이 거점으로 돌아간다.", font=font(20), fill="#384A55")
    image.save(path)


def make_map_structure(path: Path) -> None:
    image = Image.new("RGB", (1500, 760), "#F6F8FA")
    draw = ImageDraw.Draw(image)
    draw.text((60, 38), "해역 1개의 절차적 항로 구조", font=font(34, True), fill="#0B2545")
    columns = [
        ("C0", ["출항지"]),
        ("C1", ["교전", "조우", "보물"]),
        ("C2", ["항구", "교전", "정예"]),
        ("C3", ["조우", "교전", "항구"]),
        ("C4", ["교전", "정예"]),
        ("C5", ["보스"]),
    ]
    centers: list[list[tuple[int, int]]] = []
    for col_index, (label, nodes) in enumerate(columns):
        x = 120 + col_index * 245
        draw.text((x - 22, 112), label, font=font(20, True), fill="#687782")
        ys = [360] if len(nodes) == 1 else ([260, 460] if len(nodes) == 2 else [200, 360, 520])
        col_centers = []
        for y, node in zip(ys, nodes):
            color = {
                "출항지": "#D8DDD5", "교전": "#C46A51", "조우": "#7CA3A5",
                "보물": "#DDC168", "항구": "#62A979", "정예": "#DFB14F", "보스": "#BD4D4B",
            }[node]
            draw.ellipse((x - 42, y - 42, x + 42, y + 42), fill="#FFFFFF", outline=color, width=6)
            draw_text_center(draw, (x - 52, y - 34, x + 52, y + 34), node, font(18, True), "#263740")
            col_centers.append((x, y))
        centers.append(col_centers)
    for index in range(len(centers) - 1):
        for sx, sy in centers[index]:
            ordered = sorted(centers[index + 1], key=lambda point: abs(point[1] - sy))
            for tx, ty in ordered[:2]:
                draw.line((sx + 44, sy, tx - 44, ty), fill="#9BAEB7", width=3)
    draw.rectangle((55, 620, 1445, 700), fill="#FFFFFF", outline="#D8DEE5", width=2)
    draw_text_center(draw, (70, 630, 1430, 690), "각 열의 노드 유형은 셔플되며, 기본 연결 1개 + 72% 확률의 두 번째 연결을 만든 뒤 모든 노드의 진입 경로를 보정한다. 한 Act의 실제 이동은 5회다.", font(20), "#384A55")
    image.save(path)


def make_screen_layout(path: Path) -> None:
    image = Image.new("RGB", (1500, 820), "#F6F8FA")
    draw = ImageDraw.Draw(image)
    draw.text((60, 38), "플레이 화면 정보 구조", font=font(34, True), fill="#0B2545")
    draw.rectangle((70, 120, 1430, 200), fill="#0B2028", outline="#36505A", width=3)
    draw_text_center(draw, (80, 125, 620, 195), "브랜드 / 게임명", font(21, True), "#F4E5C8")
    draw_text_center(draw, (930, 125, 1420, 195), "해역 · 악명 · 전승 · 사운드", font(19, True), "#E0AE4B")
    draw.rectangle((70, 220, 1060, 630), fill="#176071", outline="#48606A", width=3)
    draw_text_center(draw, (90, 250, 1040, 600), "Canvas 1200 × 700\n항로 지도 ↔ 선박 전투", font(30, True), "#FFFFFF")
    draw.rectangle((1080, 220, 1430, 760), fill="#10252C", outline="#36505A", width=3)
    sidebar = ["선박 / 선장", "선체 · 돛 · 사기", "식량 · 식수 · 금화 · 화력", "선원", "유물", "항해일지"]
    top = 240
    for item in sidebar:
        draw.rectangle((1100, top, 1410, top + 66), fill="#172E34", outline="#36505A", width=2)
        draw_text_center(draw, (1110, top + 4, 1400, top + 62), item, font(18, True), "#F0E7D5")
        top += 78
    draw.rectangle((70, 650, 1060, 760), fill="#0D232B", outline="#36505A", width=3)
    draw_text_center(draw, (90, 660, 1040, 750), "Action Dock · 다음 항로 또는 전투 명령", font(23, True), "#F0E7D5")
    draw.text((80, 785), "880px 이하에서는 사이드 패널이 아래 2열로 이동하고, 620px 이하에서는 명령 버튼과 정보 패널이 모바일 1~2열로 재배치된다.", font=font(17), fill="#526572")
    image.save(path)


def make_resource_pipeline(path: Path) -> None:
    image = Image.new("RGB", (1500, 900), "#F6F8FA")
    draw = ImageDraw.Draw(image)
    draw.text((60, 38), "리소스 생성 파이프라인", font=font(34, True), fill="#0B2545")
    draw.text((60, 92), "별도 이미지·음원 파일 없이 소스 코드와 브라우저 API가 표시 리소스를 실시간 생성한다.", font=font(20), fill="#526572")

    lanes = [
        ("index.html", "DOM / 접근성", "HUD · 모달 · 버튼\n☠ · ♪ 기호"),
        ("styles.css", "CSS 렌더러", "자원 아이콘 · 초상\n미터 · 반응형 레이아웃"),
        ("src/game.js", "Canvas 2D", "바다 · 항로 · 선박\n포탄 · 전투 HUD"),
        ("src/game.js", "Web Audio", "발진기 기반 효과음\n20개 호출 지점"),
        ("src/game.js", "localStorage", "전승 악명 · 업그레이드\n최고 악명 · 정복 횟수"),
    ]
    y = 170
    for source, api, output in lanes:
        source_box = (60, y, 350, y + 105)
        api_box = (550, y, 900, y + 105)
        output_box = (1100, y, 1440, y + 105)
        draw.rectangle(source_box, fill="#FFFFFF", outline="#6F8794", width=3)
        draw.rectangle(api_box, fill="#E8EEF5", outline="#2E74B5", width=3)
        draw.rectangle(output_box, fill="#FFF8E8", outline="#B8862B", width=3)
        draw_text_center(draw, source_box, source, font(22, True), "#1F4D78")
        draw_text_center(draw, api_box, api, font(22, True), "#1F4D78")
        draw_text_center(draw, output_box, output, font(18, True), "#384A55")
        arrow(draw, (350, y + 52), (550, y + 52))
        arrow(draw, (900, y + 52), (1100, y + 52), color="#B8862B")
        y += 125

    draw.rectangle((60, 815, 1440, 875), fill="#FFFFFF", outline="#D8DEE5", width=2)
    draw_text_center(draw, (75, 820, 1425, 870), "배포 파일은 index.html, styles.css, src/game.js 세 계층이며 외부 CDN·fetch·미디어 프리로드가 없다.", font(19), "#384A55")
    image.save(path)


def setup_document() -> tuple[Document, int, int]:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = rgb(INK)
    normal_rpr = normal._element.get_or_add_rPr()
    normal_rpr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal_rpr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal_rpr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(BODY_AFTER)
    normal.paragraph_format.line_spacing = BODY_LINE
    normal.paragraph_format.widow_control = True

    configure_style(doc.styles["Heading 1"], 16, BLUE, True, 18, 10)
    configure_style(doc.styles["Heading 2"], 13, BLUE, True, 14, 7)
    configure_style(doc.styles["Heading 3"], 12, DARK_BLUE, True, 10, 5)
    caption = doc.styles["Caption"]
    caption.font.name = BODY_FONT
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption_rpr = caption._element.get_or_add_rPr()
    caption_rpr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run("항해의 끝: 해적왕의 유산  |  역기획서"), size=8.5, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    set_run_font(fp.add_run("NAN 2026 사전 과제  |  "), size=8.2, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    rpr.append(color)
    rpr.append(size)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    fp._p.append(fld)

    bullet_num_id = add_numbering_definition(doc, "bullet")
    decimal_num_id = add_numbering_definition(doc, "decimal")
    return doc, bullet_num_id, decimal_num_id


def build_document() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    core_loop = ASSET_DIR / "core_loop.png"
    state_flow = ASSET_DIR / "state_flow.png"
    map_structure = ASSET_DIR / "map_structure.png"
    screen_layout = ASSET_DIR / "screen_layout.png"
    resource_pipeline = ASSET_DIR / "resource_pipeline.png"
    make_core_loop(core_loop)
    make_state_flow(state_flow)
    make_map_structure(map_structure)
    make_screen_layout(screen_layout)
    make_resource_pipeline(resource_pipeline)

    doc, bullet_id, decimal_id = setup_document()

    # Editorial cover.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("BROWSER ROGUELIKE · REVERSE DESIGN DOCUMENT"), size=10.5, bold=True, color=BRASS)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("항해의 끝"), size=30, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(subtitle.add_run("해적왕의 유산"), size=21, bold=True, color=BLUE)
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_after = Pt(62)
    set_run_font(desc.add_run("구현된 플레이어블 빌드 기준 상세 역기획서"), size=13, color=MUTED)
    meta_line = doc.add_paragraph()
    meta_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_line.paragraph_format.space_after = Pt(4)
    set_run_font(meta_line.add_run("Version 1.4  |  2026-07-21"), size=10.5, bold=True, color=NAVY)
    meta_line2 = doc.add_paragraph()
    meta_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta_line2.add_run("HTML · CSS · JavaScript · Canvas 2D · LocalStorage"), size=9.5, italic=True, color=MUTED)
    statement = doc.add_paragraph()
    statement.paragraph_format.space_before = Pt(72)
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(statement.add_run("이 문서는 기획 의도를 추정하되, 수치와 동작은 현재 구현 코드를 기준으로 기록한다."), size=9.5, color=MUTED)

    page_break(doc)

    heading(doc, 1, "문서 정보")
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["문서 목적", "현재 브라우저 빌드의 규칙, 데이터, 상태 전이, UX를 재현 가능한 수준으로 명세"],
            ["분석 기준", "index.html, styles.css, src/game.js, README.md의 2026-07-21 작업 공간 버전"],
            ["제품 형태", "외부 라이브러리가 없는 정적 웹 게임 / GitHub Pages 배포 가능"],
            ["분석 원칙", "구현 사실, 설계 해석, 미구현 또는 제약을 구분"],
            ["리소스 범위", "파일형 미디어뿐 아니라 Canvas·CSS·Web Audio로 런타임 생성되는 리소스까지 포함"],
            ["범위 제외", "원본 기획서에만 있고 현재 코드에는 없는 Steam·Console 기능, 실시간 물리, 개별 선원 무장"],
        ],
        [1700, 7660],
        font_size=9.5,
        alignments=["center", "left"],
    )
    add_callout(doc, "표기 규칙", "'구현 사실'은 코드로 확인되는 동작, '설계 해석'은 동작에서 역추론한 목적, '확장 과제'는 현재 빌드에 없는 개선 방향을 뜻한다.")

    heading(doc, 2, "문서 구성")
    sections = [
        "제품 정의와 설계 축", "핵심 게임 루프", "런타임 상태와 저장 정책", "출항 준비와 초기화",
        "해역 및 절차적 항로", "자원과 경제", "선박 전투", "적 데이터와 난이도",
        "선원 및 특성", "유물과 빌드", "선택형 조우와 보물", "항구 서비스",
        "보상과 해역 진행", "전승 성장", "UI·조작·접근성", "기술 구조",
        "리소스 구성과 생성 사양", "밸런스 해석", "구현 한계와 확장 우선순위", "QA 수용 기준",
    ]
    for item in sections:
        add_number(doc, item, decimal_id)

    heading(doc, 1, "1. 제품 정의와 설계 축")
    add_callout(doc, "한 줄 정의", "한 척의 배와 한 명의 선원으로 출항해, 분기 항로의 위험과 보급을 계산하고 3개 해역의 보스를 격파하는 턴제 선박 로그라이크.", color=BRASS, fill="FFF8E8")
    add_table(
        doc,
        ["구분", "역기획 결과"],
        [
            ["장르", "로그라이크, 턴제 선박 전투, 자원 관리, 선택형 이벤트"],
            ["플랫폼", "PC 및 모바일 브라우저. 정적 파일 또는 GitHub Pages 실행"],
            ["입력", "마우스·터치 중심, 키보드 단축키 보조"],
            ["한 런의 목표", "3개 Act를 순서대로 통과하고 Act 3 보스 '심해의 왕 크라켄' 격파"],
            ["실패 조건", "선체가 0 이하이거나 사기가 0 이하"],
            ["영구 목표", "런에서 얻은 악명을 전승 자원으로 은행 처리해 3종 영구 능력을 최대 3레벨까지 강화"],
        ],
        [1800, 7560],
        alignments=["center", "left"],
    )
    heading(doc, 2, "1.1 설계 축")
    pillars = [
        ("경로의 위험 계산", "노드 유형과 현재 보급·선체 상태를 비교해 다음 1~2개 경로 중 하나를 선택한다."),
        ("전투 중 목적 전환", "선체 파괴, 돛 무력화 후 접안, 회피와 수리 중 현재 손실을 최소화하는 승리 방식을 고른다."),
        ("런 내부 시너지", "선장, 선원 역할, 특성, 유물이 포격·접안·생존·경제 빌드를 만든다."),
        ("실패의 잔존 가치", "패배해도 현재 악명은 전승 자원으로 남아 다음 런의 초기 조건을 개선한다."),
    ]
    for label, text in pillars:
        add_body(doc, f"{label}: {text}", bold_prefix=f"{label}:")

    heading(doc, 1, "2. 핵심 게임 루프")
    doc.add_picture(str(core_loop), width=Inches(6.25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_picture_alt(doc.paragraphs[-1], "핵심 게임 루프", "거점, 출항, 항로 선택, 노드 해결, 보상, 보스, 항해 종료와 전승으로 이어지는 순환 구조")
    add_caption(doc, "그림 1. 항해 단위의 핵심 루프와 전승 루프")
    heading(doc, 2, "2.1 매크로 루프")
    macro_loop_num_id = add_numbering_definition(doc, "decimal")
    for text in [
        "거점 섬에서 보유한 전승 악명을 사용해 영구 능력치를 강화한다.",
        "3명의 선장 중 하나를 선택하고 선장별 선박·초기 선원·보너스로 런을 초기화한다.",
        "각 Act의 6개 열로 구성된 항로에서 인접 노드를 선택한다.",
        "이동 시 식량과 식수를 소비하고, 도착 노드의 전투·이벤트·상점·보물을 해결한다.",
        "금화, 악명, 선원, 유물을 누적해 보스를 격파한다.",
        "승리 또는 실패 시 악명을 전승 자원으로 은행 처리하고 거점으로 돌아간다.",
    ]:
        add_number(doc, text, macro_loop_num_id)
    heading(doc, 2, "2.2 전투 마이크로 루프")
    combat_loop_num_id = add_numbering_definition(doc, "decimal")
    for text in [
        "현재 거리, 풍향, 양측 선체·돛·선원 상태를 확인한다.",
        "포격, 사슬탄, 접근, 회피, 수리, 접안, 선장 기술 중 하나를 실행한다.",
        "플레이어 행동이 유효하면 입력을 잠그고 적 턴을 처리한다.",
        "적 선체 또는 적 선원이 0이 되면 보상을 받고, 아군 선체·사기가 0이 되면 런을 종료한다.",
    ]:
        add_number(doc, text, combat_loop_num_id)

    page_break(doc)
    heading(doc, 1, "3. 런타임 상태와 저장 정책")
    doc.add_picture(str(state_flow), width=Inches(6.25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_picture_alt(doc.paragraphs[-1], "런타임 상태 전이", "interstitial, map, resolving, event, port, combat, reward, gameover, victory 상태의 전이 관계")
    add_caption(doc, "그림 2. 주요 mode 값과 화면 전이")
    add_table(
        doc,
        ["상태", "역할", "주요 이탈 조건"],
        [
            ["run 없음", "거점, 선장 선택, 전승 업그레이드", "항해 시작"],
            ["interstitial", "Act 진입 또는 정복 결과 표시", "출항·다음 해역 버튼"],
            ["map", "현재 노드와 다음 경로 표시", "경로 선택"],
            ["resolving", "이동 비용 처리와 240ms 전환 잠금", "노드 유형별 화면"],
            ["event", "2지선다 조우 또는 보물 처리", "선택 결과 후 map·combat·reward"],
            ["port", "금화로 반복 구매", "출항 버튼"],
            ["combat", "플레이어 턴과 적 턴 교대", "승리·패배"],
            ["reward", "전투 보상 또는 유물 3택", "map·다음 Act·victory"],
            ["gameover / victory", "악명 은행 처리와 결과 요약", "거점 섬"],
        ],
        [1550, 4550, 3260],
        font_size=8.8,
        alignments=["center", "left", "left"],
    )
    heading(doc, 2, "3.1 저장 범위")
    add_bullet(doc, "런 상태는 메모리에만 존재하며 새로고침 시 복구되지 않는다.", bullet_id)
    add_bullet(doc, "영구 데이터는 localStorage 키 pirate-king-legacy-meta-v1에 JSON으로 저장된다.", bullet_id)
    add_bullet(doc, "저장 항목은 전승 악명, 최고 악명, 정복 횟수, 3종 업그레이드 레벨이다.", bullet_id)
    add_bullet(doc, "저장소 사용이 실패해도 예외를 삼키고 현재 런은 계속 플레이할 수 있다.", bullet_id)
    add_callout(doc, "중요", "게임 오버와 최종 승리만 악명을 은행 처리한다. '새 항해' 메뉴로 자발적으로 포기하면 현재 악명과 유물은 소실된다.", color=RISK, fill="FBEDED")

    heading(doc, 1, "4. 출항 준비와 초기화")
    heading(doc, 2, "4.1 선장 데이터")
    add_table(
        doc,
        ["선장", "선박", "초기 보너스", "전용 기술"],
        [
            ["이사벨라 블랙배럴", "복수의 화약고", "선체 +2, 대포 +2, 포수 1명", "전탄 일제사격: 16 + floor(화력×0.8) 피해, 돛 -4, 명중 보장"],
            ["라울 스톰아이", "은빛 알바트로스", "돛 +4, 식량·식수 +3, 조타수 1명", "폭풍 가르기: 거리 3, 회피 0.8, 돛 +5"],
            ["마라 벨라돈나", "검은 세이렌", "사기 +12, 요리사 1명", "심해의 속삭임: 적 선체 -6, 돛 -7, 선원 -4, 아군 사기 +5"],
        ],
        [1700, 1500, 2500, 3660],
        font_size=8.5,
    )
    heading(doc, 2, "4.2 초기화 공식")
    add_formula(doc, "최대 선체 = 42 + 선장 선체 보너스 + 4 × 강화 용골 레벨")
    add_formula(doc, "최대 돛 = 20 + 선장 돛 보너스")
    add_formula(doc, "사기 = clamp(72 + 선장 사기 보너스 + 5 × 해적왕의 깃발 레벨, 0, 100)")
    add_formula(doc, "식량 = 식수 = 15 + 선장 보급 보너스 + 2 × 밀수업자 연줄 레벨")
    add_table(
        doc,
        ["항목", "공통 초기값"],
        [
            ["금화", "14"], ["기본 대포", "6 + 선장 대포 보너스"], ["수리도구", "2"],
            ["선원", "선장별 지정 역할 1명, 이름과 특성은 무작위"], ["유물 / 악명", "0 / 0"],
        ],
        [1900, 7460],
        alignments=["center", "left"],
    )
    add_callout(doc, "초기 체감", "전승 0레벨 기준 블랙배럴은 선체 44·표시 화력 9~12, 스톰아이는 돛 24·보급 18, 벨라돈나는 사기 84로 시작한다. 초기 선원의 등급에 따라 역할 성능이 달라져 블랙배럴의 화력도 변동한다.")

    heading(doc, 1, "5. 해역 및 절차적 항로")
    add_table(
        doc,
        ["Act", "해역", "보스", "난이도 역할"],
        [
            ["1", "평온한 시작의 바다", "붉은 산호 해적단", "기본 경제와 포격 규칙 학습"],
            ["2", "폭풍우의 중심", "왕실 철갑함 리바이어던", "적 능력치 상승과 유지비 압박"],
            ["3", "신들의 무덤", "심해의 왕 크라켄", "최고 체력 보스와 누적 빌드 검증"],
        ],
        [650, 2200, 2800, 3710],
        font_size=9.0,
        alignments=["center", "left", "left", "left"],
    )
    doc.add_picture(str(map_structure), width=Inches(6.25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_picture_alt(doc.paragraphs[-1], "절차적 항로 구조", "출항지부터 보스까지 6개 열과 13개 노드로 구성된 해역 지도")
    add_caption(doc, "그림 3. Act마다 다시 생성되는 13노드 항로")
    heading(doc, 2, "5.1 노드 배치 규칙")
    add_table(
        doc,
        ["열", "노드 수", "구성"],
        [
            ["C0", "1", "출항지"], ["C1", "3", "교전·미지의 조우·보물 1개씩, 위치 셔플"],
            ["C2", "3", "항구·교전·정예 1개씩, 위치 셔플"], ["C3", "3", "미지의 조우·교전·항구 1개씩, 위치 셔플"],
            ["C4", "2", "교전·정예 1개씩, 위치 셔플"], ["C5", "1", "보스"],
        ],
        [900, 1000, 7460],
        font_size=9.2,
        alignments=["center", "center", "left"],
    )
    add_bullet(doc, "각 노드는 다음 열에서 세로 거리가 가장 가까운 노드와 반드시 연결된다.", bullet_id)
    add_bullet(doc, "두 번째로 가까운 노드에는 72% 확률로 추가 연결을 만든다.", bullet_id)
    add_bullet(doc, "진입 연결이 없는 다음 열 노드는 가장 가까운 이전 노드에 강제로 연결해 도달 가능성을 보장한다.", bullet_id)
    add_bullet(doc, "한 경로는 각 열에서 1개 노드만 방문하므로 Act당 이동 5회, 전체 완주 시 기본 이동 15회다.", bullet_id)
    add_callout(doc, "페이싱 특성", "C4가 교전 또는 정예로 고정되고 C5가 보스이므로, 모든 Act의 끝에는 최소 2연속 전투가 발생한다.", color=BRASS, fill="FFF8E8")

    heading(doc, 2, "5.2 노드 유형")
    add_table(
        doc,
        ["유형", "표식", "기능", "직접 보상 또는 위험"],
        [
            ["교전", "전", "일반 적과 전투", "금화·악명"], ["정예", "정", "Act별 정예함과 전투", "금화·악명·유물 3택"],
            ["미지의 조우", "?", "5종 이벤트 중 균등 무작위", "자원·선원·정예 전투 등"], ["항구", "항", "금화로 반복 구매", "보급·수리·고용·사기·수리도구"],
            ["숨겨진 보물", "$", "위험한 유물 또는 안전한 금화 선택", "유물 / 금화 18·악명 3"], ["해역 지배자", "왕", "Act 보스 전투", "금화·악명·유물, 다음 Act"],
        ],
        [1300, 600, 3200, 4260],
        font_size=8.8,
        alignments=["center", "center", "left", "left"],
    )

    page_break(doc)
    heading(doc, 1, "6. 자원과 경제")
    add_table(
        doc,
        ["자원", "획득", "소모", "0 또는 상한 처리"],
        [
            ["선체", "수리, 심해 강철판", "적 포격·접안·폭풍·금고 함정", "0 이하 즉시 게임 오버"],
            ["돛", "항구 수리, 응급수리, 폭풍 가르기", "사슬탄 피격·폭풍", "0이면 접근·회피 불가"],
            ["사기", "구조, 선술집, Act 정복, 특성·유물", "굶주림·갈증·피격·이벤트", "0 이하 반란으로 게임 오버, 상한 100"],
            ["식량", "초기 보급, 항구, 구조 이벤트, Act 정복", "노드 이동·이벤트", "이동 후 0이면 사기 -10"],
            ["식수", "초기 보급, 항구, 구조 이벤트, Act 정복", "노드 이동·이벤트", "이동 후 0이면 사기 -12"],
            ["금화", "전투·보물·세이렌", "항구·검문 통행세", "하한 별도 없음, 부족하면 구매 버튼 비활성"],
            ["악명", "전투·이벤트·유물", "런 중 소모 없음", "종료 시 전승 악명으로 전환"],
        ],
        [1000, 2700, 2800, 2860],
        font_size=8.25,
    )
    heading(doc, 2, "6.1 이동 비용")
    add_formula(doc, "기본 보급 비용 = max(1, 2 - [스톰아이 선장] - [요리사 1명 이상])")
    add_formula(doc, "실제 식량 비용 = 기본 보급 비용 + [주당 특성 1명 이상]")
    add_formula(doc, "실제 식수 비용 = 기본 보급 비용")
    add_formula(doc, "이동 사기 회복 = [노래하는 선수상] + Σ(요리사 등급 숙련 단계)")
    add_bullet(doc, "스톰아이와 요리사의 감소 효과는 함께 적용되지만 하한이 1이므로 둘을 동시에 보유해도 비용은 0이 되지 않는다.", bullet_id)
    add_bullet(doc, "노래하는 선수상과 레어 이상 요리사의 사기 회복은 합산해 이동마다 먼저 적용한다.", bullet_id)
    add_bullet(doc, "식량과 식수가 모두 0이 된 이동은 총 사기 -22를 유발한다.", bullet_id)

    heading(doc, 2, "6.2 금화 보정")
    add_formula(doc, "최종 금화 = round(기본 금화 × (1 + 0.30×[보물지도] + 0.10×[행운아 존재]))")
    add_callout(doc, "중첩 규칙", "행운아는 여러 명이어도 +10%만 적용되고, 도금된 보물지도와는 가산되어 최대 1.4배가 된다.")

    heading(doc, 1, "7. 선박 전투")
    add_callout(doc, "전투 성격", "실시간 포물선 물리 대신, 거리 1~3과 풍향을 명중률에 반영하는 턴제 명령 전투다. 행동 1회 뒤 적 행동 1회가 처리된다.", color=BRASS, fill="FFF8E8")
    heading(doc, 2, "7.1 전투 초기화")
    add_bullet(doc, "일반·정예 전투는 거리 2 또는 3에서 균등 무작위 시작하며, 보스는 거리 3에서 시작한다.", bullet_id)
    add_bullet(doc, "풍향은 순풍·측풍·역풍 중 균등 무작위, 풍속은 1~3이다.", bullet_id)
    add_bullet(doc, "선장 기술은 전투마다 1회 충전되며 다음 전투에서 다시 사용할 수 있다.", bullet_id)
    add_bullet(doc, "풍속 값은 현재 화면 표시에만 사용되고 공식에는 참여하지 않는다.", bullet_id)

    heading(doc, 2, "7.2 명중률과 피해 공식")
    add_formula(doc, "선체 포격 명중률 = clamp(0.84 - 0.09×(거리-1) + 풍향 + 선장 + 유물, 0.25, 0.96)")
    add_formula(doc, "사슬탄 명중률 = clamp(0.76 - 0.11×(거리-1) + 풍향 + 선장 + 유물, 0.25, 0.96)")
    add_table(
        doc,
        ["보정", "값", "적용 조건"],
        [
            ["순풍", "+0.06", "두 탄종"], ["역풍", "-0.10", "병 속의 폭풍이 없을 때"],
            ["블랙배럴", "+0.06", "선장 고유 상시 보정"], ["황동 육분의", "+0.10", "유물 보유"],
        ],
        [1800, 1200, 6360],
        alignments=["center", "center", "left"],
    )
    add_formula(doc, "선체 포격 피해 = 표시 화력 + randomInt(2,6) + 3×[왕실 흑색화약]")
    add_formula(doc, "사슬탄 돛 피해 = randomInt(6,10) + Σ(포수 1 + 등급 숙련 단계)")
    add_formula(doc, "표시 화력 = 런의 대포 수 + Σ(포수 1 + 등급 숙련 단계)")

    heading(doc, 2, "7.3 플레이어 명령")
    add_table(
        doc,
        ["명령", "사용 조건", "효과", "리스크"],
        [
            ["선체 포격", "항상", "적 선체 피해, 25% 확률 적 선원 -1", "거리·풍향에 따른 빗나감"],
            ["사슬탄", "항상", "적 돛 6~10 + 등급별 포수 보너스", "선체를 직접 줄이지 않음"],
            ["접근 기동", "거리 >1, 아군 돛 >0", "성공 시 거리 -1, 회피 0.08", "실패해도 턴 소모"],
            ["회피 기동", "아군 돛 >0", "거리 +1 또는 거리 3에서 회피 0.28", "직접 피해 없음"],
            ["응급수리", "도구 >0, 선체가 최대 미만", "선체 7+Σ(수리공 3+2×숙련), 돛 +3", "수리도구 1 소모"],
            ["접안 공격", "거리 1, 적 돛 ≤55%", "성공 시 즉시 나포 승리", "실패 시 선체 5~9, 사기 -6, 25% 선원 손실"],
            ["선장 기술", "전투당 1회", "선장별 고유 강효과", "재사용 불가"],
        ],
        [1200, 1900, 3450, 2810],
        font_size=8.25,
    )
    heading(doc, 2, "7.4 기동과 접안 공식")
    add_formula(doc, "접근 성공률 = clamp(0.68 + Σ(조타수 0.10+0.05×숙련) + 0.12×(현재 돛/최대 돛), 0.30, 0.96)")
    add_formula(doc, "접안 성공률 = clamp(0.42 + (아군 전투력 - 적 선원)/45, 0.20, 0.90)")

    heading(doc, 2, "7.5 적 턴 우선순위")
    enemy_turn_num_id = add_numbering_definition(doc, "decimal")
    for text in [
        "거리 3이고 적 돛이 남아 있으면 58% 확률로 거리 2까지 접근한다.",
        "그 외 거리 1이고 적 선원이 아군 전투력의 80%를 넘으면 30% 확률로 적 접안 공격을 수행한다.",
        "그 외에는 포격 명중 판정을 수행한다.",
        "턴 종료 후 28% 확률로 풍향과 풍속을 다시 뽑는다.",
    ]:
        add_number(doc, text, enemy_turn_num_id)
    add_formula(doc, "적 명중률 = clamp(0.73 - 0.07×거리 - 회피 + 0.03×ActIndex, 0.10, 0.88)")
    add_formula(doc, "적 선체 피해 = 적 damage + randomInt(0,4)")
    add_bullet(doc, "명중하면 사기 -2, 28% 확률로 돛 3~6 추가 피해가 발생한다.", bullet_id)
    add_bullet(doc, "적 접안은 선체 4~7 + ActIndex 피해와 사기 -5를 준다.", bullet_id)

    page_break(doc)
    heading(doc, 1, "8. 적 데이터와 난이도 스케일링")
    add_table(
        doc,
        ["분류", "Act", "선체", "돛", "선원", "damage", "기본 보상: 금화 / 악명"],
        [
            ["일반", "1", "25~31", "15", "10", "5", "11 / 7"],
            ["일반", "2", "34~40", "17", "13", "6", "16 / 9"],
            ["일반", "3", "43~49", "19", "16", "7", "21 / 11"],
            ["정예", "1", "40", "20", "16", "7", "20 / 15"],
            ["정예", "2", "52", "23", "20", "8", "27 / 19"],
            ["정예", "3", "64", "26", "24", "9", "34 / 23"],
            ["보스", "1", "54", "24", "20", "8", "28 / 28"],
            ["보스", "2", "69", "28", "25", "10", "38 / 36"],
            ["보스", "3", "84", "32", "30", "12", "48 / 44"],
        ],
        [950, 500, 850, 700, 700, 760, 4900],
        font_size=8.4,
        alignments=["center", "center", "center", "center", "center", "center", "center"],
    )
    add_bullet(doc, "일반 적 이름은 Act별 3종 풀에서 무작위 선택된다.", bullet_id)
    add_bullet(doc, "정예와 보스는 Act별 고정 이름을 사용한다.", bullet_id)
    add_bullet(doc, "정예와 보스는 승리 후 유물 3택을 제공하며, 일반 적은 금화·악명만 제공한다.", bullet_id)
    add_bullet(doc, "나포 성공 시 기본 금화에 +8을 더한 뒤 금화 배율을 적용하고 악명 +5를 추가한다.", bullet_id)
    heading(doc, 2, "8.1 승리 및 보상 판정")
    add_formula(doc, "적 선체 ≤ 0 또는 적 선원 ≤ 0 → 승리")
    add_formula(doc, "승리 사기 회복 = 4×[끝없는 럼주통] + 2×[굳센 특성 존재]")

    heading(doc, 1, "9. 선원 및 특성")
    add_body(doc, "선원은 이름 10종, 역할 5종, 특성 4종을 조합해 생성된다. 선장 선택 시 1명이 주어지고, 구조 이벤트 또는 항구 고용으로 최대 4명까지 확장한다.")
    heading(doc, 2, "9.1 선원 객체 데이터 구조")
    add_table(
        doc,
        ["필드", "형식·예시", "생성 규칙과 사용처"],
        [
            ["id", "문자열 · 172…-0.483…", "Date.now()와 Math.random()을 결합한 런타임 식별자"],
            ["name", "문자열 · 애꾸눈 모건", "CREW_NAMES 10종에서 균등 무작위 선택"],
            ["roleId", "gunner 등 영문 키", "역할 조회, 인원수 계산, 전투·항해 공식에 사용"],
            ["role", "문자열 · 포수", "역할의 한글 표시명"],
            ["mark", "문자열 · 포", "선원 목록의 28×28 아바타에 표시"],
            ["basePower / power", "정수 · 1~4 / 1~8", "역할 기본값 / 기본값에 등급 보너스를 더한 실제 전투력"],
            ["effect", "문자열", "등급이 반영된 역할 효과 설명. 선원 행과 hover title에 사용"],
            ["trait", "객체 · id/name/effect", "TRAITS 4종 중 하나를 통째로 참조"],
            ["rarityId / rarity", "영문 키 / 한글명", "normal·rare·epic·legendary 조회와 UI 표시"],
            ["rarityBonus", "정수 · 0/1/2/4", "등급에 따른 power 가산값"],
            ["crewTier", "정수 · 0/1/2/3", "역할 효과 강화 단계"],
        ],
        [1300, 3000, 5060],
        font_size=8.25,
        alignments=["center", "left", "left"],
    )
    add_formula(doc, "일반 영입 선원 = 이름 10종 × 역할 5종 × 특성 4종 × 등급 4종 = 800개 명목 조합")
    add_formula(doc, "선장 초기 선원 = 이름 10종 × 선장 고정 역할 1종 × 특성 4종 × 등급 4종 = 선장당 160개 명목 조합")

    heading(doc, 2, "9.2 이름 데이터")
    add_table(
        doc,
        ["번호", "이름", "번호", "이름"],
        [
            ["1", "애꾸눈 모건", "6", "무쇠턱 로사"],
            ["2", "붉은 수염 앤", "7", "까마귀 핀"],
            ["3", "도끼손 잭", "8", "조용한 벤"],
            ["4", "썰물의 니아", "9", "파도칼 미라"],
            ["5", "북극성 톰", "10", "노을빛 산초"],
        ],
        [900, 3780, 900, 3780],
        font_size=9.0,
        alignments=["center", "left", "center", "left"],
    )
    add_bullet(doc, "이름 선택 확률은 각각 10%이며, 이미 보유한 이름을 제외하는 검사가 없어 동일 이름 선원이 중복될 수 있다.", bullet_id)

    heading(doc, 2, "9.3 역할 데이터")
    add_table(
        doc,
        ["roleId / 역할 / mark", "기본 power", "노말 기준 역할 효과", "등급 강화"],
        [
            ["gunner / 포수 / 포", "3", "표시 화력·포격·사슬탄 +1", "숙련 단계마다 +1"],
            ["rigger / 조타수 / 타", "2", "접근 성공률 +10%", "숙련 단계마다 +5%p"],
            ["carpenter / 수리공 / 수", "2", "응급수리 선체 +3", "숙련 단계마다 +2"],
            ["cook / 요리사 / 요", "1", "이동 보급 비용 -1", "숙련 단계만큼 이동 사기 회복"],
            ["marine / 갑판전사 / 전", "4", "접안 전투력 +4", "숙련 단계마다 +2"],
        ],
        [2300, 1100, 3100, 2860],
        font_size=8.2,
        alignments=["center", "center", "left", "left"],
    )

    heading(doc, 2, "9.4 특성 데이터")
    add_table(
        doc,
        ["trait id / 특성", "효과", "중첩 방식"],
        [
            ["steady / 굳센", "전투 승리 시 사기 +2", "존재 여부만 확인, 다수 비중첩"],
            ["drunk / 주당", "선술집 회복 +5, 이동 식량 소모 +1", "존재 여부만 확인, 다수 비중첩"],
            ["lucky / 행운아", "금화 획득 +10%", "존재 여부만 확인, 다수 비중첩"],
            ["scarred / 상처투성이", "접안 전투력 +2", "인원수만큼 중첩"],
        ],
        [2300, 3800, 3260],
        font_size=8.6,
        alignments=["center", "left", "left"],
    )
    add_bullet(doc, "특성 선택 확률은 각각 25%이며, 역할·이름과 독립적으로 복원 추출한다.", bullet_id)

    heading(doc, 2, "9.5 선원 등급 데이터")
    add_table(
        doc,
        ["rarityId / 등급", "기본 획득 확률", "power 보너스", "역할 숙련 단계", "UI 색상"],
        [
            ["normal / 노말", "55%", "+0", "0", "#AEB9B6"],
            ["rare / 레어", "28%", "+1", "1", "#54A9C2"],
            ["epic / 에픽", "13%", "+2", "2", "#B982D9"],
            ["legendary / 레전드", "4%", "+4", "3", "#FFD36F"],
        ],
        [2200, 1800, 1700, 1800, 1860],
        font_size=8.8,
        alignments=["center", "center", "center", "center", "center"],
    )
    add_bullet(doc, "등급은 이름·역할·특성과 독립적으로 매 생성 시 복원 추출하며, 선장 초기 선원에도 같은 확률을 적용한다.", bullet_id)
    add_formula(doc, "선원 실제 power = 역할 basePower + 등급 rarityBonus")
    add_table(
        doc,
        ["역할", "노말", "레어", "에픽", "레전드"],
        [
            ["포수", "피해 +1", "피해 +2", "피해 +3", "피해 +4"],
            ["조타수", "접근 +10%", "접근 +15%", "접근 +20%", "접근 +25%"],
            ["수리공", "선체 +3", "선체 +5", "선체 +7", "선체 +9"],
            ["요리사", "보급 -1", "보급 -1·사기 +1", "보급 -1·사기 +2", "보급 -1·사기 +3"],
            ["갑판전사", "접안 +4", "접안 +6", "접안 +8", "접안 +10"],
        ],
        [1500, 1900, 1900, 1900, 2160],
        font_size=8.1,
        alignments=["center", "center", "center", "center", "center"],
    )

    heading(doc, 2, "9.6 선장별 초기 선원")
    add_table(
        doc,
        ["선장", "고정 roleId", "초기 선원에서 무작위인 값", "기획 효과"],
        [
            ["이사벨라 블랙배럴", "gunner / 포수", "이름, 특성, 등급", "포격 빌드의 최소 작동 조건을 보장"],
            ["라울 스톰아이", "rigger / 조타수", "이름, 특성, 등급", "접근 기동 성공률을 기본 보완"],
            ["마라 벨라돈나", "cook / 요리사", "이름, 특성, 등급", "장기 항해의 보급 비용을 기본 완화"],
        ],
        [2100, 2000, 2260, 3000],
        font_size=8.6,
        alignments=["center", "center", "center", "left"],
    )

    heading(doc, 2, "9.7 선원 전투력")
    add_formula(doc, "전투력 = 8 + Σ(역할 basePower + 등급 bonus) + Σ(갑판전사 4+2×숙련) + 2×상처투성이 수 + 5×[크라켄의 이빨]")
    add_callout(doc, "해석", "갑판전사 1명의 총 전투력 기여는 노말 +8, 레어 +11, 에픽 +14, 레전드 +18이다. 역할 기본 power·등급 power 보너스·접안 역할 보너스를 모두 합산한다.")

    heading(doc, 2, "9.8 획득·손실·표시 규칙")
    add_table(
        doc,
        ["시점", "조건·비용", "데이터 변화", "제약·후속 처리"],
        [
            ["항해 시작", "선장 선택", "선장 고정 역할의 선원 1명 생성", "이름·특성·등급은 무작위, 선원 수 1로 시작"],
            ["구명정 구조", "현재 선원 <4", "이름·역할·특성·등급 무작위 선원 +1, 사기 +5", "4명이면 선택 비활성화"],
            ["항구 고용", "금화 ≥14, 현재 선원 <4", "금화 -14, 이름·역할·특성·등급 무작위 선원 +1", "같은 항구에서 금화가 허용하는 만큼 반복 가능"],
            ["직접 내보내기", "항로 화면, 현재 선원 >1", "확인 후 지정 선원 1명 제거", "보상 없음, 마지막 1명은 버튼 비활성화"],
            ["접안 실패", "선원 >1이며 25% 판정 성공", "무작위 배열 인덱스의 선원 1명 제거", "이 규칙만으로는 마지막 1명을 잃지 않음"],
            ["런 종료·포기", "승리·패배·포기", "run 객체와 함께 선원 목록 폐기", "선원은 전승 저장 대상이 아님"],
            ["선원 패널", "항상", "mark, name, role·trait, 등급, +power 표시", "등급색 테두리·badge와 역할·특성 hover 설명"],
        ],
        [1600, 2200, 3260, 2300],
        font_size=7.9,
        alignments=["center", "left", "left", "left"],
    )
    add_callout(doc, "현재 미구현", "선원 순서 변경·장비·레벨·체력·부상·영구 보존은 없다. 이름·역할·특성·등급 중복도 허용된다.", color=BRASS, fill="FFF8E8")

    heading(doc, 1, "10. 유물과 빌드")
    add_body(doc, "유물 획득 시 등급을 가중 추첨한 뒤 해당 등급의 미보유 유물 중 최대 3개를 제시한다. 한 런의 보유 한도는 6개이며 모든 유물은 고유하다. 획득할 때마다 악명 +3을 얻고, 한도가 찼거나 후보가 없으면 유물을 악명 8로 자동 변환한다.")
    add_table(
        doc,
        ["유물", "등급", "효과", "주요 빌드"],
        [
            ["황동 육분의", "노말", "포격·사슬탄 명중률 +10%", "안정 포격"],
            ["왕실 흑색화약", "레전드", "선체 포격 피해 +3, 블랙배럴 기술 피해 +3", "고화력 포격"],
            ["노래하는 선수상", "노말", "노드 이동마다 사기 +1", "장기 생존"],
            ["크라켄의 이빨", "에픽", "접안 전투력 +5", "나포"],
            ["병 속의 폭풍", "노말", "역풍 명중률 -10% 제거", "기상 안정"],
            ["도금된 보물지도", "레어", "금화 획득 +30%", "경제"],
            ["심해 강철판", "레어", "최대 선체 +8 및 획득 즉시 선체 +8", "방어"],
            ["끝없는 럼주통", "에픽", "전투 승리 시 사기 +4", "연전"],
        ],
        [2100, 1200, 3960, 2100],
        font_size=8.35,
    )
    heading(doc, 2, "10.1 대표 시너지")
    add_bullet(doc, "포격: 블랙배럴 + 포수 + 황동 육분의 + 왕실 흑색화약.", bullet_id)
    add_bullet(doc, "나포: 갑판전사 + 상처투성이 + 크라켄의 이빨 + 사슬탄.", bullet_id)
    add_bullet(doc, "항해 생존: 스톰아이 또는 요리사 + 노래하는 선수상 + 끝없는 럼주통.", bullet_id)
    add_bullet(doc, "경제: 행운아 + 도금된 보물지도 + 항구 경유.", bullet_id)

    heading(doc, 2, "10.2 등급 추첨·관리 규칙")
    add_table(
        doc,
        ["등급", "후보 1칸 기본 확률", "포함 유물"],
        [
            ["노말", "55%", "황동 육분의, 노래하는 선수상, 병 속의 폭풍"],
            ["레어", "28%", "도금된 보물지도, 심해 강철판"],
            ["에픽", "13%", "크라켄의 이빨, 끝없는 럼주통"],
            ["레전드", "4%", "왕실 흑색화약"],
        ],
        [1500, 2300, 5560],
        font_size=8.6,
        alignments=["center", "center", "left"],
    )
    add_bullet(doc, "3택은 복원 없는 순차 추출이다. 이미 제시한 유물과 보유 유물을 제외하고, 남은 유물이 있는 등급만 대상으로 기본 가중치를 재정규화한다.", bullet_id)
    add_bullet(doc, "유물 버리기는 모달이 닫힌 항로 화면에서만 가능하다. 확인 후 효과가 즉시 사라지며 보상은 없고, 이후 같은 유물을 다시 획득할 수 있다.", bullet_id)
    add_bullet(doc, "심해 강철판을 버리면 최대 선체가 8 감소하고 현재 선체가 새 최대치를 넘을 경우 함께 보정된다.", bullet_id)
    add_bullet(doc, "보상 화면에서 보유 한도 6개가 이미 찬 경우에는 버릴 유물을 고르는 교체 단계 없이 악명 8로 자동 변환한다.", bullet_id)

    heading(doc, 1, "11. 선택형 조우와 보물")
    add_body(doc, "미지의 조우 노드에 진입하면 5개 이벤트 중 하나를 균등 무작위로 선택한다. 대부분 두 선택지를 제공하며, 비용을 지불할 수 없는 선택지는 비활성화된다.")
    add_table(
        doc,
        ["이벤트", "선택", "조건", "결과"],
        [
            ["부서진 구명정", "생존자 구조", "선원 <4", "무작위 선원 +1, 사기 +5"],
            ["", "보급 상자", "없음", "식량·식수 +5, 사기 -3"],
            ["수평선을 삼킨 폭풍", "정면 돌파", "없음", "선체 4~10 피해, 돛 2~5 피해, 악명 +8. 스톰아이는 선체 2~5"],
            ["", "우회", "없음", "식량·식수 -3, 둘 중 하나가 0이면 사기 -8"],
            ["세이렌의 암초", "근원 탐색", "없음", "55% 성공: 금화 16·악명 +6 / 실패: 사기 -16. 벨라돈나는 확정 성공"],
            ["", "지나간다", "없음", "사기 -2"],
            ["왕실 검문선", "통행세", "금화 ≥10", "금화 -10 후 복귀"],
            ["", "해적기", "없음", "즉시 정예 전투"],
            ["선창의 불온한 속삭임", "비상 식량", "식량 ≥4", "식량 -4, 사기 +14. 주당 보유 시 +19"],
            ["", "주동자 처벌", "없음", "50% 확률 사기 +8, 아니면 -12"],
        ],
        [1700, 1550, 1400, 4710],
        font_size=7.8,
    )
    heading(doc, 2, "11.1 숨겨진 보물")
    add_table(
        doc,
        ["선택", "결과"],
        [
            ["금고 해체", "35% 확률로 선체 4~9 피해 후 유물 3택. 피해로 선체가 0이 되면 즉시 종료"],
            ["금화만 회수", "조정 전 금화 18과 악명 +3. 금화는 보물지도·행운아 배율 적용"],
        ],
        [2000, 7360],
        font_size=9.2,
    )

    heading(doc, 1, "12. 항구 서비스")
    add_body(doc, "항구는 출항 전까지 같은 서비스를 반복 구매할 수 있는 유일한 안전 경제 노드다. 구매 후 화면을 다시 그려 현재 금화와 조건을 즉시 반영한다.")
    add_table(
        doc,
        ["서비스", "가격", "효과", "비활성 조건"],
        [
            ["보급품 묶음", "8", "식량 +6, 식수 +6", "금화 부족"],
            ["선체 수리", "10", "선체 +12, 돛 +5", "금화 부족 또는 둘 다 최대"],
            ["선원 고용", "14", "무작위 선원 +1", "금화 부족 또는 선원 4명"],
            ["선술집", "6", "사기 +15, 주당 보유 시 +20", "금화 부족 또는 사기 100"],
            ["수리도구", "7", "전투용 수리도구 +1", "금화 부족"],
        ],
        [1900, 900, 3600, 2960],
        font_size=8.8,
        alignments=["center", "center", "left", "left"],
    )
    add_callout(doc, "경제 선택", "전투 1회의 기본 금화 보상은 Act 1 일반 기준 11이다. 따라서 초반에는 수리 1회 또는 보급 1회가 일반 전투 1회의 수입 대부분을 사용한다.", color=BRASS, fill="FFF8E8")

    heading(doc, 1, "13. 보상과 해역 진행")
    heading(doc, 2, "13.1 전투 보상")
    add_formula(doc, "전투 금화 = adjustedGold(기본 금화 + 8×[나포])")
    add_formula(doc, "전투 악명 = 기본 악명 + 5×[나포]")
    add_bullet(doc, "일반 전투는 보상 확인 후 지도에 복귀한다.", bullet_id)
    add_bullet(doc, "정예 전투는 보상 확인 후 유물 3택을 거쳐 지도에 복귀한다.", bullet_id)
    add_bullet(doc, "보스 전투는 유물 3택 후 Act를 완료한다.", bullet_id)
    heading(doc, 2, "13.2 Act 완료")
    add_bullet(doc, "Act 1·2 완료 시 식량 +4, 식수 +4, 사기 +8을 회복한다.", bullet_id)
    add_bullet(doc, "다음 Act는 새로운 항로를 생성하고 현재 선체·돛·금화·선원·유물은 유지한다.", bullet_id)
    add_bullet(doc, "Act 3 보스 격파 후에는 해적왕의 탄생 승리 화면으로 전환한다.", bullet_id)

    heading(doc, 1, "14. 전승 성장")
    add_body(doc, "런 종료 시 현재 악명을 전승 악명에 전부 더한다. 동시에 최고 악명을 갱신하고, 승리한 경우 정복 횟수를 1 증가시킨다. 영구 강화는 각 3레벨까지 구매할 수 있다.")
    add_table(
        doc,
        ["업그레이드", "레벨당 효과", "Lv1 / Lv2 / Lv3 비용", "누적 비용", "최대 효과"],
        [
            ["강화 용골", "초기 최대 선체 +4", "18 / 38 / 64", "120", "+12"],
            ["밀수업자 연줄", "초기 식량·식수 +2", "16 / 34 / 58", "108", "각 +6"],
            ["해적왕의 깃발", "초기 사기 +5", "20 / 40 / 68", "128", "+15"],
        ],
        [2100, 2100, 2300, 1300, 1560],
        font_size=8.8,
        alignments=["center", "left", "center", "center", "center"],
    )
    add_callout(doc, "완전 해금 비용", "3종 업그레이드를 모두 최대화하려면 전승 악명 356이 필요하다.")

    heading(doc, 1, "15. UI·조작·접근성")
    doc.add_picture(str(screen_layout), width=Inches(6.25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_picture_alt(doc.paragraphs[-1], "플레이 화면 정보 구조", "상단 상태, 중앙 캔버스, 하단 명령, 우측 선박 패널의 화면 배치")
    add_caption(doc, "그림 4. 데스크톱 기준 화면 영역과 반응형 재배치")
    heading(doc, 2, "15.1 정보 우선순위")
    add_bullet(doc, "상단: 현재 해역, 런 악명, 전승 악명, 사운드 토글.", bullet_id)
    add_bullet(doc, "중앙 Canvas: 항로와 전투 상황을 상호 배타적으로 시각화.", bullet_id)
    add_bullet(doc, "하단 Action Dock: 현재 상태에서 가능한 선택만 버튼으로 제공.", bullet_id)
    add_bullet(doc, "우측 Ship Panel: 선체·돛·사기, 보급, 금화, 화력, 등급 badge와 관리 버튼이 있는 선원·유물, 최근 로그 8개.", bullet_id)
    add_bullet(doc, "Modal Layer: 선장 선택, 이벤트, 항구, 보상, 결과를 입력 차단형으로 표시.", bullet_id)
    heading(doc, 2, "15.2 반응형 규칙")
    add_table(
        doc,
        ["구간", "레이아웃 변화"],
        [
            [">1120px", "메인 화면 + 330px 우측 패널"],
            ["≤1120px", "우측 패널 288px, Canvas 최소 높이 380px"],
            ["≤880px", "게임과 패널을 1열로 전환, 패널 내부는 2열"],
            ["≤620px", "상단 정보 축약, 명령 2열, 선원·유물·일지는 전체 폭, 모달 1열"],
        ],
        [1800, 7560],
        alignments=["center", "left"],
    )
    heading(doc, 2, "15.3 입력 매핑")
    add_table(
        doc,
        ["맥락", "마우스·터치", "키보드"],
        [
            ["항로", "노드 또는 Action Dock 버튼", "숫자 1~3"],
            ["선원·유물 관리", "항로 화면의 내보내기·버리기 버튼과 확인 모달", "직접 단축키 없음"],
            ["모달 선택", "선택 카드", "숫자 1~9, Enter로 주요 버튼"],
            ["전투", "명령 버튼", "F 포격, S 사슬탄, A 접근, D 회피, R 수리, B 접안, Q 기술"],
            ["사운드", "음표 아이콘", "직접 단축키 없음"],
        ],
        [1600, 3300, 4460],
        font_size=8.8,
    )
    heading(doc, 2, "15.4 피드백")
    add_bullet(doc, "선체·돛·사기는 수치와 색상 미터를 동시에 제공한다.", bullet_id)
    add_bullet(doc, "선원·유물은 노말·레어·에픽·레전드 등급을 색상 테두리와 badge로 함께 표시한다.", bullet_id)
    add_bullet(doc, "불가능한 행동은 disabled 상태로 비활성화한다.", bullet_id)
    add_bullet(doc, "전투 행동 중 locked를 사용해 적 턴 처리 전 중복 입력을 차단한다.", bullet_id)
    add_bullet(doc, "Web Audio API의 짧은 발진음으로 포격·승리·피해·선택 피드백을 제공한다.", bullet_id)
    add_bullet(doc, "prefers-reduced-motion 환경에서는 CSS 전환 시간을 사실상 제거한다.", bullet_id)

    heading(doc, 1, "16. 기술 구조")
    add_table(
        doc,
        ["계층", "구현"],
        [
            ["문서 구조", "index.html의 상단 상태, Canvas, Action Dock, Ship Panel, Modal Layer"],
            ["표현", "styles.css의 데스크톱·태블릿·모바일 반응형 규칙"],
            ["게임 로직", "src/game.js 단일 모듈의 데이터 상수, 상태 객체, 전이 함수, 렌더 함수"],
            ["그래픽", "Canvas 2D API로 파도, 항로, 노드, 선박, 포탄을 매 프레임 직접 그림"],
            ["UI 갱신", "DOM 요소를 캐시하고 상태 변화 시 HUD·버튼·모달을 재구성"],
            ["오디오", "AudioContext 발진기 기반의 짧은 효과음"],
            ["영구 저장", "localStorage JSON"],
            ["배포", "빌드 단계 없는 정적 파일, GitHub Actions Pages 워크플로"],
        ],
        [1800, 7560],
        font_size=9.2,
        alignments=["center", "left"],
    )
    heading(doc, 2, "16.1 핵심 런 데이터")
    add_body(doc, "런 객체는 captainId, actIndex, mode, map, currentNodeId, hull/maxHull, sails/maxSails, morale, food, water, gold, cannons, repairKits, infamy, crew, artifacts, logs, combat, travelCount, banked를 보유한다.")
    heading(doc, 2, "16.2 구현상 장점")
    add_bullet(doc, "서버나 패키지 설치 없이 브라우저에서 즉시 실행된다.", bullet_id)
    add_bullet(doc, "게임 데이터가 상수 배열로 모여 있어 선장·유물·경제 수치를 빠르게 조정할 수 있다.", bullet_id)
    add_bullet(doc, "Canvas와 DOM UI를 분리해 연출과 정보 전달의 책임이 명확하다.", bullet_id)
    add_bullet(doc, "모든 외부 상태가 localStorage 하나에 제한되어 GitHub Pages에 적합하다.", bullet_id)

    heading(doc, 1, "17. 리소스 구성과 생성 사양")
    doc.add_picture(str(resource_pipeline), width=Inches(6.25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_picture_alt(doc.paragraphs[-1], "리소스 생성 파이프라인", "HTML, CSS, JavaScript가 DOM, CSS, Canvas 2D, Web Audio, localStorage를 거쳐 게임 표시 리소스를 생성하는 구조")
    add_caption(doc, "그림 5. 정적 소스에서 런타임 표시·사운드·저장 리소스로 이어지는 흐름")
    add_callout(doc, "핵심 결론", "현재 제출 빌드는 이미지·SVG·스프라이트·음원·웹폰트·외부 라이브러리 파일을 사용하지 않는다. 플레이에 필요한 리소스는 전체 소스에 포함된 도형, 색상, 문자열, 수치로 생성된다.", color=BRASS, fill="FFF8E8")

    heading(doc, 2, "17.1 배포 리소스 인벤토리")
    add_table(
        doc,
        ["분류", "실제 파일·API", "역할 및 포함 관계"],
        [
            ["문서", "index.html", "게임 영역, HUD, 모달, 접근성 레이블, 해골·음표 문자 기호를 포함"],
            ["스타일", "styles.css", "팔레트, 레이아웃, 미터, CSS 자원 아이콘, 선장 초상, 반응형 규칙을 포함"],
            ["로직·생성", "src/game.js", "게임 데이터, Canvas 렌더링, 효과음 합성, 상태 전이, 입력 처리를 포함"],
            ["이미지·벡터 파일", "0개", "PNG, JPG, WebP, GIF, SVG, 스프라이트시트 참조 없음"],
            ["오디오 파일", "0개", "MP3, WAV, OGG 참조 없음. Web Audio로 런타임 합성"],
            ["번들 글꼴", "0개", "WOFF, TTF, OTF를 배포하지 않고 운영체제 글꼴 스택 사용"],
            ["외부 의존성", "0개", "CDN, 패키지, fetch 요청 없이 브라우저 기본 API만 사용"],
        ],
        [1700, 2500, 5160],
        font_size=8.7,
        alignments=["center", "center", "left"],
    )
    add_callout(doc, "제출 관점", "전체 소스 3계층이 곧 전체 리소스다. 별도 assets 폴더가 없어 GitHub Pages에 그대로 올려도 누락될 미디어 경로가 없다.")

    heading(doc, 2, "17.2 시각 리소스 사양")
    add_table(
        doc,
        ["리소스", "생성 방식", "구성 요소와 변형"],
        [
            ["바다·하늘", "Canvas 2D", "단색 하늘·심해·해수 면과 12행 사인파 포말. 전투 장면에는 반투명 구름 5개 추가"],
            ["항로", "Canvas 2D", "베지어 곡선 연결선, 방문·선택 상태별 점선 패턴, 선택 가능 노드의 펄스 링"],
            ["항로 노드", "Canvas 2D", "원형 배지와 한글·기호 마크. 출·전·정·?·항·$·왕의 7종"],
            ["선박", "Canvas 2D", "선체 다각형, 포문 5개, 돛, 돛대, 깃발. 좌우 반전·크기·선장 외투색으로 변형"],
            ["섬 실루엣", "Canvas 2D", "출항 전 화면의 산형 다각형과 해안 직사각형"],
            ["전투 HUD", "Canvas 2D", "양측 선체·돛 게이지, 풍향·풍속, 거리, 선원 전투력, 턴 표시"],
            ["포탄 효과", "Canvas 2D", "원형 포탄이 sin(progress×π) 포물선을 따라 이동. 양측 발사 방향과 색상 분리"],
            ["자원 아이콘", "CSS 도형", "식량, 물방울, 금화, 대포 4종. border, border-radius, pseudo-element로 생성"],
            ["선장 초상", "CSS 도형", "48×48 프레임 안 얼굴 원과 외투 사다리꼴. 선장별 portrait·coat 색상 적용"],
            ["브랜드·조작 기호", "문자 글리프", "해골 ☠, 사운드 ♪, 풍속 ▸, 선원 역할·항로 유형의 한글 한 글자 마크"],
        ],
        [1700, 2500, 5160],
        font_size=8.15,
        alignments=["center", "center", "left"],
    )

    heading(doc, 2, "17.3 해역별 Canvas 팔레트")
    add_table(
        doc,
        ["해역", "하늘 sky", "바다 sea", "심해 deep", "포말 foam", "강조 accent"],
        [
            ["Act 1 시작의 바다", "#9EC4C0", "#176071", "#0B3F51", "#B8D9D2", "#E0AE4B"],
            ["Act 2 폭풍우의 중심", "#697D83", "#184451", "#0B2B36", "#9BB6B7", "#D98B58"],
            ["Act 3 신들의 무덤", "#342F49", "#193F4C", "#101C2A", "#8DA8A0", "#C77262"],
        ],
        [1300, 1612, 1612, 1612, 1612, 1612],
        font_size=8.25,
        alignments=["center", "center", "center", "center", "center", "center"],
    )
    add_body(doc, "공통 UI 팔레트는 심해색 #071A24, 패널 #10252C·#172E34, 금색 #E0AE4B·#FFD36F, 위험색 #C54D3F, 회복색 #57AE79, 돛색 #54A9C2를 중심으로 구성된다. 등급 토큰은 노말 #AEB9B6, 레어 #54A9C2, 에픽 #B982D9, 레전드 #FFD36F를 사용한다.")

    heading(doc, 2, "17.4 효과음 큐시트")
    add_table(
        doc,
        ["상황", "주파수·파형", "길이·음량", "의도"],
        [
            ["선장 선택 / 전승 구매 / 사운드 ON", "420 / 620 / 520Hz · triangle", "0.05 / 0.10 / 0.07초 · 기본", "짧고 밝은 UI 확인음"],
            ["교전 시작", "160Hz · sawtooth", "0.16초 · 0.045", "저역 경고와 긴장 전환"],
            ["포격 명중 / 빗나감 / 사슬탄", "92Hz square / 130Hz sine / 110Hz square", "0.18 / 0.08 / 0.14초", "중량감, 공백감, 금속 충격을 구분"],
            ["접근 / 회피 / 수리", "260 / 330 / 500Hz · triangle", "각 0.07 / 0.07 / 0.09초", "기동 방향과 회복 성공을 음높이로 구분"],
            ["접안 성공 / 실패", "680 / 82Hz · sawtooth", "0.18초 0.040 / 0.20초 0.050", "극단적인 고저 차이로 결과 전달"],
            ["블랙배럴 / 스톰아이 / 벨라돈나 기술", "76 square / 720 triangle / 190 sine", "0.28 / 0.15 / 0.30초", "중포, 고속 기동, 주술의 캐릭터성 표현"],
            ["적 접안 / 적 포격 명중", "84Hz sawtooth / 72Hz square", "0.18초 0.045 / 0.20초 0.055", "피격 경고와 저역 충격"],
            ["전투 승리 / 유물 획득", "540 / 760Hz · triangle", "0.16초 0.050 / 0.15초 0.045", "보상 단계의 상승감"],
            ["최종 승리 / 패배", "720Hz triangle / 72Hz sawtooth", "0.35 / 0.45초 · 0.055", "런 종료 결과를 가장 긴 음으로 강조"],
        ],
        [1900, 3000, 1700, 2760],
        font_size=7.65,
        alignments=["center", "left", "center", "left"],
    )
    add_bullet(doc, "playTone은 OscillatorNode와 GainNode를 매 호출 생성하고, 지정 시간 뒤 oscillator.stop으로 종료한다.", bullet_id)
    add_bullet(doc, "음소거 상태에서는 생성을 생략하며, AudioContext 생성이 실패하면 음소거로 전환해 게임 진행을 유지한다.", bullet_id)
    add_bullet(doc, "BGM, 루프 앰비언스, 음량 단계 조절은 현재 구현에 없다.", bullet_id)

    heading(doc, 2, "17.5 글꼴과 문자 리소스")
    add_table(
        doc,
        ["용도", "글꼴 스택", "배포·대체 정책"],
        [
            ["일반 UI", "Pretendard → Noto Sans KR → Apple SD Gothic Neo → system-ui → sans-serif", "웹폰트를 포함하지 않으며 설치된 첫 글꼴을 사용"],
            ["제목·강조", "Georgia → Noto Serif KR → serif", "영문은 Georgia, 한글은 사용 가능한 세리프 대체 글꼴로 표시"],
            ["Canvas 텍스트", "Georgia 또는 system-ui → sans-serif", "맵 제목·전투 HUD·노드 라벨을 브라우저 글꼴로 즉시 렌더링"],
            ["문자 아이콘", "☠, ♪, ▸, $, ?, 출·전·정·항·왕", "이미지 대체 없이 Unicode·ASCII 글리프 사용. 운영체제별 형태 차이 가능"],
        ],
        [1700, 3500, 4160],
        font_size=8.25,
        alignments=["center", "left", "left"],
    )
    add_callout(doc, "라이선스 범위", "게임 저장소에 제3자 미디어와 글꼴 파일을 재배포하지 않는다. 운영체제 글꼴은 사용자 환경에서 참조하고, 시각·사운드 리소스 정의는 프로젝트 소스 코드에 포함된다.")

    heading(doc, 2, "17.6 로딩·성능·교체 영향")
    add_table(
        doc,
        ["항목", "현재 동작", "기획·운영 영향"],
        [
            ["초기 요청", "index.html, styles.css, src/game.js", "미디어 다운로드와 디코딩 단계가 없어 별도 로딩 화면이 필요하지 않음"],
            ["캐시 무효화", "CSS·JS URL에 ?v=20260721e", "배포 갱신 시 쿼리 버전을 변경해야 브라우저 캐시 혼선을 줄일 수 있음"],
            ["프레임 렌더", "1200×700 Canvas를 requestAnimationFrame으로 지속 갱신", "파도·선박 흔들림·선택 노드 펄스·포탄 궤적을 매 프레임 다시 그림"],
            ["오디오 초기화", "최초 효과음 시 AudioContext 생성", "브라우저 자동재생 정책상 사용자 입력 이후 정상 활성화되는 구조"],
            ["영구 데이터", "pirate-king-legacy-meta-v1 단일 JSON", "이미지·오디오 캐시와 무관하며 저장 용량이 매우 작음"],
            ["향후 파일 리소스 교체", "현재 프리로더·리소스 매니저 없음", "스프라이트·BGM을 추가하면 로딩 상태, 오류 대체, 저작권 표기, 용량 예산이 새로 필요"],
        ],
        [1900, 3000, 4460],
        font_size=8.15,
        alignments=["center", "left", "left"],
    )

    heading(doc, 1, "18. 밸런스 해석")
    heading(doc, 2, "18.1 선장별 초반 경제")
    add_table(
        doc,
        ["선장", "초기 핵심값", "기본 이동비", "예상 플레이"],
        [
            ["블랙배럴", "선체 44, 표시 화력 9~12", "식량 2 / 식수 2", "포수 등급에 따라 초반 화력이 달라지는 공격형"],
            ["스톰아이", "돛 24, 식량·식수 18", "식량 1 / 식수 1", "항로 유지와 회피를 통해 장기 손실 최소화"],
            ["벨라돈나", "사기 84, 요리사", "식량 1 / 식수 1", "이벤트 안정성과 사기 여유를 이용한 선택형"],
        ],
        [1800, 2400, 1800, 3360],
        font_size=8.7,
    )
    add_callout(doc, "리스크 곡선", "전승 0레벨 블랙배럴은 완주 기본 이동 15회에 식량·식수 각각 30이 필요해 초기 15만으로는 부족하다. 반면 스톰아이와 요리사를 가진 벨라돈나는 기본 15로 정확히 15회를 감당한다.", color=BRASS, fill="FFF8E8")
    heading(doc, 2, "18.2 의사결정 압력")
    add_bullet(doc, "정예는 유물을 보장하지만 전투 손실과 수리비를 높인다.", bullet_id)
    add_bullet(doc, "사슬탄은 즉시 승리 속도를 늦추지만 접안 보너스와 나포 보상을 연다.", bullet_id)
    add_bullet(doc, "항구에서 보급·수리·선원을 동시에 살 수 없어 현재 런의 병목 자원을 우선해야 한다.", bullet_id)
    add_bullet(doc, "C4 전투 뒤 곧바로 보스가 오므로 수리도구와 선장 기술 관리가 Act 후반 생존을 좌우한다.", bullet_id)
    heading(doc, 2, "18.3 확률과 변동성")
    add_bullet(doc, "맵 유형 총량은 고정되고 위치만 섞여 경로 생성의 변동성은 제한적이다.", bullet_id)
    add_bullet(doc, "명중·이벤트·선원 특성·유물 3택이 실제 런 차이를 만든다.", bullet_id)
    add_bullet(doc, "최소·최대 clamp가 극단적 조합에서도 완전 확정 또는 완전 불가능 판정을 방지한다.", bullet_id)
    add_bullet(doc, "접안은 최대 90%이므로 강한 빌드도 실패 리스크를 유지한다.", bullet_id)

    heading(doc, 1, "19. 구현 한계와 확장 우선순위")
    add_callout(doc, "범위 구분", "아래 항목은 현재 구현의 결함이라기보다, 원본 콘셉트와 완성형 제품 사이의 차이를 역기획 관점에서 정리한 것이다.")
    add_table(
        doc,
        ["우선", "현재 구현", "영향", "확장 방향"],
        [
            ["P0", "풍속 1~3이 표시 전용", "전략 정보와 실제 규칙 불일치", "명중·기동·포탄 피해 중 하나에 풍속 계수 연결"],
            ["P0", "돛 피해가 단계적으로 적 AI와 포격을 약화하지 않음", "사슬탄 가치가 접안 조건에 편중", "돛 비율에 따른 기동·명중·턴 우선권 페널티"],
            ["P0", "런 저장·이어하기 없음", "모바일 이탈 시 런 손실", "버전이 있는 runState 직렬화와 복구"],
            ["P1", "항로 구조와 유형 총량이 Act마다 동일", "재플레이 맵 다양성 제한", "Act별 열 수, 노드 풀, 잠금 경로, 위험 지형 추가"],
            ["P1", "선원은 역할·특성만 있고 체력·무장·배치가 없음", "백병전 깊이 제한", "선원 슬롯, 부상, 무기, 배치 보너스"],
            ["P1", "전승은 3종 수치 강화만 존재", "장기 해금 동기 제한", "선장·선박·유물 풀 해금과 도전 과제"],
            ["P1", "적 AI가 풍향과 자신의 손상도를 거의 고려하지 않음", "전투 패턴 반복", "상태 기반 전술 프로필과 보스 전용 패턴"],
            ["P2", "특성 중첩 규칙이 UI에 노출되지 않음", "빌드 기대와 실제 계산 불일치", "툴팁에 '고유/중첩' 명시"],
            ["P2", "시드·통계·튜토리얼 없음", "재현 QA와 초기 학습 어려움", "seed 입력, 런 리포트, 첫 항해 온보딩"],
        ],
        [650, 2750, 2500, 3460],
        font_size=7.7,
    )
    heading(doc, 2, "19.1 확장 시 유지해야 할 정체성")
    add_bullet(doc, "경로 선택이 자원 소비와 직접 연결되는 구조.", bullet_id)
    add_bullet(doc, "선체 파괴와 나포가 서로 다른 보상 곡선을 만드는 전투 목표.", bullet_id)
    add_bullet(doc, "선장·선원·유물의 짧고 읽기 쉬운 시너지.", bullet_id)
    add_bullet(doc, "패배도 다음 항해의 출발점을 바꾸는 전승 루프.", bullet_id)

    heading(doc, 1, "20. QA 수용 기준")
    heading(doc, 2, "20.1 출항과 저장")
    for text in [
        "세 선장 모두 설명과 일치하는 선박·자원·초기 선원을 받는다.",
        "업그레이드 구매 시 비용이 차감되고 새 런 초기값에 반영된다.",
        "패배·승리 후 악명이 정확히 한 번만 은행 처리된다.",
        "자발적 포기 시 현재 런 악명이 전승되지 않는다.",
    ]:
        add_bullet(doc, text, bullet_id)
    heading(doc, 2, "20.2 항로와 자원")
    for text in [
        "각 Act에 13개 노드가 생성되고 보스까지 도달 가능한 연결이 존재한다.",
        "다음 연결 노드 이외의 Canvas 클릭은 이동을 발생시키지 않는다.",
        "이동 비용이 선장·요리사·주당 조합에 맞게 계산된다.",
        "선체 또는 사기가 0이 되는 즉시 다른 선택보다 게임 오버가 우선한다.",
    ]:
        add_bullet(doc, text, bullet_id)
    heading(doc, 2, "20.3 전투")
    for text in [
        "모든 행동 버튼이 조건과 일치하게 활성·비활성화된다.",
        "행동 직후 적 턴 종료 전 추가 입력이 처리되지 않는다.",
        "유물·선원·선장 보정이 표시 명중률과 실제 판정에 동일하게 반영된다.",
        "접안 성공 시 나포 추가 보상, 실패 시 피해·사기·선원 손실 판정이 적용된다.",
        "정예·보스 승리 후 유물 선택 흐름이 끊기지 않는다.",
    ]:
        add_bullet(doc, text, bullet_id)
    heading(doc, 2, "20.4 등급과 보유 관리")
    for text in [
        "선원·유물 기본 등급 추첨 가중치가 노말 55%, 레어 28%, 에픽 13%, 레전드 4%로 합계 100%를 이룬다.",
        "선원 power가 역할 basePower와 등급 보너스 +0/+1/+2/+4를 합산해 표시 전투력에 반영된다.",
        "선원 역할 숙련 단계 0/1/2/3이 포격·접근·수리·이동 사기·접안 성능에 등급별 수치로 반영된다.",
        "유물 3택에 동일 유물이 중복되지 않고, 미보유 유물이 없는 등급은 남은 가중치로 재정규화된다.",
        "항로 이외 상태에서는 내보내기·버리기가 비활성화되고, 마지막 선원 1명은 직접 제거할 수 없다.",
        "선원·유물 제거는 확인 모달을 거치며 취소 시 데이터가 변하지 않고 확정 시 보상 없이 즉시 제거된다.",
        "심해 강철판을 버릴 때 최대 선체 -8과 현재 선체 상한 보정이 적용된다.",
    ]:
        add_bullet(doc, text, bullet_id)
    heading(doc, 2, "20.5 UI와 호환성")
    for text in [
        "마우스·터치와 키보드가 같은 행동 결과를 만든다.",
        "620px 이하에서 명령·자원·모달 텍스트가 잘리거나 겹치지 않는다.",
        "사운드가 차단된 브라우저에서도 게임 로직은 정상 진행된다.",
        "localStorage를 사용할 수 없는 환경에서도 런은 플레이 가능하다.",
    ]:
        add_bullet(doc, text, bullet_id)

    heading(doc, 1, "부록 A. 구현 사실 요약")
    add_table(
        doc,
        ["지표", "값"],
        [
            ["선장 / 선원 이름 / 역할 / 특성 / 등급", "3 / 10 / 5 / 4 / 4"],
            ["선원 시작 인원 / 보유 한도", "1 / 4"],
            ["선원·유물 등급 확률", "노말 55% / 레어 28% / 에픽 13% / 레전드 4%"],
            ["선원 등급 성능", "power +0/+1/+2/+4, 역할 숙련 0/1/2/3"],
            ["유물", "8종(노말 3 / 레어 2 / 에픽 2 / 레전드 1), 최대 6개 보유"],
            ["보유 관리", "항로 화면에서 선원 내보내기 / 유물 버리기, 확인 모달, 보상 없음"],
            ["해역 / 보스", "3 / 3"],
            ["해역당 전체 노드 / 실제 이동", "13 / 5"],
            ["이벤트 / 보물 이벤트", "5 / 1"],
            ["항구 서비스", "5"],
            ["전투 명령", "7"],
            ["전승 업그레이드", "3종 × 3레벨"],
            ["Canvas", "1200 × 700"],
            ["정적 배포 계층", "index.html / styles.css / src/game.js"],
            ["외부 이미지 / 오디오 / 웹폰트 / 라이브러리", "0 / 0 / 0 / 0"],
            ["합성 효과음 호출 지점", "20"],
            ["영구 저장 키", "pirate-king-legacy-meta-v1"],
            ["분석 소스", "index.html · styles.css · src/game.js · README.md"],
        ],
        [3600, 5760],
        font_size=8.5,
        alignments=["left", "left"],
    )

    # Metadata and core properties.
    doc.core_properties.title = "항해의 끝: 해적왕의 유산 - 상세 역기획서"
    doc.core_properties.subject = "NAN 2026 해커톤 사전 과제 브라우저 게임 역기획"
    doc.core_properties.author = "NAN 2026 Hackathon Team"
    doc.core_properties.keywords = "로그라이크, 해적, 선박 전투, 역기획서, 브라우저 게임"
    doc.core_properties.comments = "현재 구현 소스, 선원·유물 등급 데이터, 런타임 생성 리소스 기준 Version 1.4"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
